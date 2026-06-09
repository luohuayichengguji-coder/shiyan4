from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs"
FINAL_DOCX = OUT / "AI开发技术-实验4.1-医学图像技术调研报告.docx"
FIG = OUT / "figures"
TABLES = OUT / "tables"

CN_FONT = "宋体"
EN_FONT = "Times New Roman"


def set_run_font(run, size=None, bold=None, color=None, cn_font=CN_FONT, en_font=EN_FONT):
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn("w:ascii"), en_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), en_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def paragraph(text="", size=10.5, bold=False, align=None, before=0, after=0, line=1.5, first_line=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line:
        pf.first_line_indent = Pt(21)
    if align is not None:
        p.alignment = align
        if align == WD_ALIGN_PARAGRAPH.CENTER:
            pf.first_line_indent = None
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    return p


def heading(text, level=1):
    size = 12 if level <= 2 else 11
    p = paragraph(text, size=size, bold=True, before=4 if level == 1 else 2, after=0, first_line=False)
    return p


def add_blank(lines=1):
    for _ in range(lines):
        paragraph("", first_line=False, line=1.0)


def add_cover():
    add_blank(6)
    p = paragraph("《AI开发技术》实验报告", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, line=1.0)
    add_blank(10)
    paragraph("医学图像技术调研报告", size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, line=1.0)
    add_blank(8)
    for line in [
        "学    号：    2302320122    ",
        "姓    名：      胡文韬      ",
        "班    级：   大数据12301    ",
        "指导教师：       黄晋       ",
        "报告日期：2026 年 6月 1  日",
    ]:
        p = paragraph(line, size=16, bold=True, first_line=False)
        p.paragraph_format.left_indent = Inches(1.35)
    doc.add_page_break()


def read_csv(name):
    path = TABLES / name
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def set_cell_text(cell, text, size=9, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(dxa))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)


def add_table(headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, size=font_size, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], "EDEDED")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], val, size=font_size, align=align)
    paragraph("", size=4, first_line=False, line=1.0)
    return table


def add_picture(path, caption):
    if not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(5.75))
    cp = paragraph(caption, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, first_line=False, line=1.0)
    return cp


def create_numbering(pattern="%1）"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, attrs in [
        ("w:start", {"w:val": "1"}),
        ("w:numFmt", {"w:val": "decimal"}),
        ("w:lvlText", {"w:val": pattern}),
        ("w:lvlJc", {"w:val": "left"}),
    ]:
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        lvl.append(el)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "420")
    ind.set(qn("w:hanging"), "300")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(p, num_id):
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def numbered_items(items):
    num_id = create_numbering()
    for item in items:
        p = paragraph("", first_line=False)
        apply_num(p, num_id)
        r = p.add_run(item)
        set_run_font(r, size=10.5)


def setup_document():
    global doc
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = EN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("AI开发技术-实验4.1 医学图像技术调研")
    set_run_font(r, size=9)


def truncate(text, n=68):
    text = " ".join(str(text).split())
    return text if len(text) <= n else text[: n - 1] + "…"


setup_document()
add_cover()

heading("1.实验题目及要求")
heading("1.1实验题目", level=2)
numbered_items([
    "围绕医学图像处理和眼底视网膜病变分割开展技术调研，理解眼底病灶分割在糖尿病辅助诊断中的意义。",
    "熟悉 Segment Anything Model（SAM）及其医学图像适配方法，分析其迁移到眼底视网膜病变分割领域的可行性。",
    "结合本地论文库与最新科研论文，梳理 SOTA 方法、常用数据集、评价指标、研究不足和可改进方向。",
    "使用 Codex 桌面端、Agent 与多个调研/文档 Skill 完成任务拆解、文献分析、方案设计、实验规划、Word 报告和 PPT。",
])

heading("1.2实验内容", level=2)
numbered_items([
    "根据提供论文分类，整理 SAM 在医学图像分割中的应用，并将其扩展到眼底视网膜病变/糖尿病视网膜病变病灶分割领域。",
    "使用 Codex 桌面端等 AI 工具创建项目，选中现有论文文件夹，对 29 篇本地 PDF 进行题名、年份、方向、数据集和指标线索抽取。",
    "进一步分析本地论文内容，在线检索 2024-2026 年相关最新科研论文，比较通用 SAM、医学 SAM 适配、眼底 DR 病灶分割、低质量眼底增强和 SAM 蒸馏方法。",
    "围绕“迁移到特定医学图像分割，特别是眼底视网膜分割”和“待分割图像质量较低”两个问题，提出 LQ-Fundus-SAM 研究方案。",
    "根据上述问题给出实验规划，并输出图文并茂的技术调研报告、调研内容 PPT、合规矩阵和 AI 环境/Skill 调用轨迹。",
])

heading("1.3重点调研问题", level=2)
numbered_items([
    "当前医学图像分割和眼底 DR 病灶分割领域的 SOTA 方法有哪些？它们使用哪些数据集和评价指标？",
    "通用 SAM 迁移到眼底病灶分割时，为什么会在小病灶、低对比边界和类别不平衡场景下出现不稳定？",
    "MedSAM、Medical SAM Adapter、SAM-Med2D、TP-DRSeg 等方法对本课题有哪些可借鉴点？",
    "如果构建一个面向低质量眼底图像的 SAM 适配方案，baseline、模块设计、损失函数、消融实验和风险备选应如何规划？",
])

heading("2.实验环境配置")
heading("2.1硬件环境", level=2)
numbered_items([
    "调研生成主环境：项目记录为 conda hj 环境，工作目录 /mnt/c/Users/VictorTau/Desktop/hj-5。",
    "本次报告重写与渲染环境：macOS 14.6，Darwin arm64，工作目录 /Users/victorhu/Desktop/hj-5。",
    "说明：本实验以文献调研、文档生成和方案规划为主，不包含真实模型训练，因此不编造 GPU 型号、训练耗时或实验分数。",
])

heading("2.2软件环境", level=2)
numbered_items([
    "Python：主调研记录为 Python 3.12.13；本次 DOCX 重写使用 Codex workspace dependencies Python。",
    "AI 辅助工具：Codex 桌面端，结合 Agent/Skill 工作流完成任务拆解、资料整理和文档生成。",
    "主要文档工具：python-docx 生成 Word，python-pptx 生成 PPT，LibreOffice + PyMuPDF/render_docx.py 用于渲染检查。",
    "输出目录：outputs/ 保存最终报告、PPT、阶段性记录、表格、图表和渲染检查结果；used_skills_md/ 保存 Skill 审批材料。",
])

heading("2.3核心依赖库", level=2)
add_table(
    ["模块类型", "依赖包名称", "用途说明"],
    [
        ["PDF 读取", "PyMuPDF / fitz", "抽取本地论文题名、页数、摘要/引言片段、arXiv/DOI、数据集和指标线索"],
        ["表格处理", "pandas / csv", "生成本地论文索引、重点论文分析、SOTA、数据集和指标对比表"],
        ["报告生成", "python-docx", "生成可直接提交的 Word 技术调研报告"],
        ["PPT 生成", "python-pptx", "生成课程汇报 PPTX，并检查幻灯片数量和可打开性"],
        ["图表绘制", "matplotlib", "绘制技术路线图、问题-方案图、实验流程图等 PNG/SVG/PDF 图表"],
        ["渲染检查", "LibreOffice / render_docx.py", "将 DOCX 转换为页面 PNG，检查版式、分页、图片和表格是否正常"],
    ],
    [1900, 2050, 5280],
    font_size=8.5,
)

heading("2.4Codex + Agent + Skill配置过程", level=2)
numbered_items([
    "读取任务指导文件 AI技术-实验4.1-实验要求-医学图像技术调研.txt，明确实验目的、内容、步骤、要求和相关资料。",
    "建立 outputs/ 目录，用脚本抽取本地 PDF，生成 local_paper_index.csv/json/md 与 local_pdf_extracted_snippets.md。",
    "读取并参考 academic-research-suite、nature-reader、nature-academic-search、technical-report-writer、nature-writing、nature-figure、documents、presentations 等 Skill。",
    "nature-academic-search 批量检索接口曾出现 asyncio 事件循环错误，随后降级为网页线索 + 单篇 arXiv/DOI 核验，并在过程记录中说明。",
    "最终把环境、命令、Agent 任务拆解、Skill 调用和降级处理写入 outputs/AI_environment_and_skill_trace.md。",
])

heading("3阶段一：提示词方式完成医学图像技术调研")
heading("3.1 阶段一任务说明", level=2)
paragraph("阶段一按照任务指导文件中的提示词提纲进行：先问 SOTA 方法、评估方式和数据集，再读 PDF 想 idea，分析领域问题和相似领域迁移思路，最后选择 baseline 和改进方案，生成图文并茂的技术调研报告。该阶段更接近“提示词驱动的线性调研”，重点是快速形成研究问题和初步方案。")

heading("3.2提示词内容", level=2)
numbered_items([
    "调研提示：现在这个领域的 SOTA 方法有哪些？使用的评估方式、数据集有哪些？请详细分析这些 SOTA 方法。",
    "读 PDF 想 idea 提示：当前领域存在哪些问题？眼底图像分割与裂缝检测、骨架检测、边缘检测等相似领域有哪些可迁移思路？",
    "解决问题提示：分析需要使用哪个模型作为 baseline 比较合适，如何在该 baseline 上实施改进方法？",
    "报告生成提示：依据技术调研报告提纲，生成图文并茂的实验报告，要求保存为 Word 格式，可以直接提交。",
])

heading("3.3 生成结果", level=2)
paragraph("阶段一生成了本地论文库索引、重点论文初读记录、SOTA 方法族地图、数据集和指标表，并形成“通用 SAM → 医学 SAM 适配 → 眼底专用病灶分割 → 低质量增强 → 轻量化蒸馏”的主线。")
local_rows = read_csv("local_paper_index.csv")
category_counts = {}
for row in local_rows:
    category_counts[row["category"]] = category_counts.get(row["category"], 0) + 1
add_table(
    ["论文类别", "篇数", "调研作用"],
    [
        ["SAM模型蒸馏", category_counts.get("SAM模型蒸馏", 0), "分析 MobileSAM、EdgeSAM、KD-SAM 等轻量化思路"],
        ["SAM相关", category_counts.get("SAM相关", 0), "分析 SAM、MedSAM、SAM-Med2D、Medical SAM Adapter 等基础与医学适配模型"],
        ["医学图像增强", category_counts.get("医学图像增强", 0), "分析 SAT-Net、AMIR、DiffCode 等低质量眼底增强方法"],
        ["眼底视网膜病变分割", category_counts.get("眼底视网膜病变分割", 0), "分析 RTNet、GlanceSeg、TP-DRSeg 等 DR 病灶分割方法"],
    ],
    [1900, 900, 6460],
    font_size=8.5,
)

add_picture(FIG / "technical_roadmap.png", "图1 LQ-Fundus-SAM 技术路线图")

heading("3.4阶段一遇到的问题", level=2)
numbered_items([
    "问题描述：通用 SAM 文献多以自然图像为主，直接迁移到眼底图像时缺少微动脉瘤、出血、渗出等病灶先验。解决方式：引入 MedSAM、Medical SAM Adapter、SAM-Med2D 和 TP-DRSeg 等医学/眼底适配论文作为核心依据。",
    "问题描述：数据集信息容易混淆，APTOS、EyePACS、Messidor 等常用于分类或筛查，不应误写为像素级病灶分割主数据。解决方式：区分 IDRiD、DDR、FGADR 等分割数据和分类数据，只把后者作为预训练或泛化讨论。",
    "问题描述：在线检索会遇到预印本时间和接口核验问题。解决方式：记录检索日期、来源、筛选理由，对当前日期之后的条目仅作为待跟踪线索。",
    "问题描述：本实验是技术调研与实验规划，没有真实训练结果。解决方式：报告中只提出待验证假设，不编造 Dice、IoU 或 AUC 数值提升。",
])

heading("3.5调研全过程记录", level=2)
numbered_items([
    "读取任务指导文件，拆解实验目的、实验内容、实验步骤和最终交付物要求。",
    "用 PyMuPDF 扫描本地 29 篇 PDF，生成论文索引、摘要片段和数据集/指标线索。",
    "选择 SAM、MedSAM、Medical SAM Adapter、SAM-Med2D、RTNet、GlanceSeg、SAT-Net、KD-SAM 等重点论文进行结构化阅读。",
    "围绕 2024-2026 年医学 SAM、DR 病灶分割、低质量眼底增强和 SAM 蒸馏进行补充检索。",
    "归纳研究不足，提出 LQ-Fundus-SAM 方案，并给出数据集、baseline、损失函数、消融实验和可视化规划。",
])

heading("3.6阶段一成果文件", level=2)
numbered_items([
    "outputs/tables/local_paper_index.csv：本地 29 篇 PDF 论文索引。",
    "outputs/notes/key_paper_analysis.md：重点论文人工分析表。",
    "outputs/notes/literature_search_record.md：在线检索记录、来源、筛选理由和核验方式。",
    "outputs/notes/sota_datasets_metrics.md：SOTA 方法、数据集和指标总结。",
    "outputs/notes/research_problem_solution.md：研究不足和 LQ-Fundus-SAM 方案设计。",
    "outputs/notes/experiment_plan.md：实验数据、预处理、训练、对比、消融、风险和备选方案。",
])

heading("3.7阶段一调研结论", level=2)
paragraph("阶段一证明，仅依靠单次提示词可以较快形成调研框架，但容易出现证据链不够细、论文来源不够可追溯、数据集属性混淆等问题。对于医学图像技术调研，必须把本地 PDF、在线检索、表格证据、图表和最终报告串成可核验流程。")

heading("4. 阶段二：Agent + 多 Skill 协同完成技术调研报告")
heading("4.1阶段二任务说明", level=2)
paragraph("阶段二采用 Agent + 多 Skill 协同模式，将“医学图像技术调研报告”拆分为任务理解、论文索引、重点论文阅读、在线检索、SOTA 综合、问题分析、方案设计、实验规划、报告生成、PPT 生成和渲染检查等子任务。每个子任务都有明确产物，便于复查和继续迭代。")

heading("4.2配置过程", level=2)
numbered_items([
    "在 Codex 桌面端打开 hj-5 工作目录，优先读取 AGENTS.md 和任务指导文件。",
    "确认 conda hj 环境与 Python 依赖可用，制图时将 MPLCONFIGDIR 指向 outputs/.mplconfig，避免默认目录不可写。",
    "创建 outputs/scripts/，把 PDF 抽取、重点论文笔记、研究产物构建、图表绘制、报告/PPT 构建和渲染检查脚本集中保存。",
    "创建 used_skills_md/，保存本次使用或参考的 Skill 摘要和原始 Markdown，便于老师审核。",
])

heading("4.3 Skill体系构建", level=2)
skill_rows = [
    ["Skill 1", "academic-research-suite", "阶段化研究拆解、问题建模、实验规划"],
    ["Skill 2", "nature-reader", "本地重点论文结构化阅读，不做全文翻译"],
    ["Skill 3", "nature-academic-search", "在线检索和 arXiv/DOI 单篇核验"],
    ["Skill 4", "technical-report-writer", "报告结构、证据映射、合规矩阵"],
    ["Skill 5", "nature-writing", "研究现状、问题-方案链条组织"],
    ["Skill 6", "nature-figure", "技术路线图、问题-方案图、实验流程图"],
    ["Skill 7", "documents / presentations", "Word 报告、PPTX 生成与渲染检查"],
]
add_table(["编号", "Skill名称", "作用"], skill_rows, [900, 2600, 5860], font_size=8.5)

heading("4.4Agent 任务拆解轨迹", level=2)
numbered_items([
    "读取任务指导文件，确认报告必须覆盖研究背景、现状、存在问题、研究方案、实验规划和参考文献。",
    "扫描本地论文文件夹，将论文分为 SAM 相关、医学 SAM、眼底病变分割、医学图像增强、SAM 蒸馏等方向。",
    "抽取重点论文摘要、方法、实验和结论片段，再人工归纳研究问题、方法、数据集、指标、优势、不足和可借鉴点。",
    "在线检索最新论文，筛选 TP-DRSeg、Medical SAM 2、MedSAM2、KD-SAM、SAT-Net 等与本课题最相关的工作。",
    "构建 SOTA 方法族、数据集和评价指标对比，避免只按论文顺序堆叠摘要。",
    "提出 LQ-Fundus-SAM，并规划 baseline、模块、损失函数、消融实验、可视化和风险备选方案。",
    "生成最终 Word 报告、PPTX、合规矩阵、AI 环境与 Skill 调用轨迹，并进行渲染检查。",
])

heading("4.5 Agent 构建过程与 Skill 调用顺序", level=2)
paragraph("Agent 的调用顺序从“证据收集”逐步推进到“方案生成”和“交付物检查”：先用文献阅读与检索 Skill 建立证据，再用写作、制图、文档和演示 Skill 形成可提交产物。")
numbered_items([
    "调用本地 PDF 抽取脚本，生成 29 篇论文索引和摘要片段。",
    "调用重点论文分析流程，形成 key_paper_analysis.csv 和 key_paper_analysis.md。",
    "调用在线检索流程，对 2024-2026 年相关论文建立检索记录。",
    "调用 SOTA 综合流程，生成方法族、数据集和指标对比。",
    "调用方案设计流程，形成 LQ-Fundus-SAM 的模块设计和总损失函数。",
    "调用文档与演示流程，输出 DOCX、PPTX、README、合规矩阵和渲染检查图。",
])

heading("4.6项目文件", level=2)
numbered_items([
    "outputs/AI开发技术-实验4.1-医学图像技术调研报告.docx：最终 Word 实验报告。",
    "outputs/AI开发技术-实验4.1-医学图像技术调研汇报.pptx：最终调研汇报 PPT。",
    "outputs/AI开发技术-实验4.1-医学图像技术调研报告.md：报告 Markdown 中间稿。",
    "outputs/compliance_matrix.md：对照任务指导文件的逐条合规矩阵。",
    "outputs/AI_environment_and_skill_trace.md：AI 环境配置、命令、Skill/Agent 调用轨迹和降级处理记录。",
    "outputs/figures/：技术路线图、SOTA 方法族地图、数据集指标矩阵、问题方案映射、实验流程图。",
    "used_skills_md/：本次使用或参考的 Skill 原始 Markdown 与审批摘要。",
])

heading("4.7系统优化与技术要点分析", level=2)
paragraph("本次调研不是实现一个训练完成的模型，而是从论文证据中提炼可执行研究方案。核心技术要点包括医学 SAM 领域适配、低质量眼底增强、小病灶/边界约束和轻量化蒸馏。")
add_picture(FIG / "sota_method_map.png", "图2 SOTA 方法族地图")

sota_rows = read_csv("sota_method_comparison.csv")
add_table(
    ["方向", "代表方法", "核心思想", "主要局限"],
    [[truncate(r["方向"], 16), truncate(r["代表方法"], 28), truncate(r["核心思想"], 42), truncate(r["主要局限"], 38)] for r in sota_rows],
    [1500, 2450, 3000, 2410],
    font_size=7.5,
)

heading("4.7.1 重点论文分析", level=3)
key_rows = read_csv("key_paper_analysis.csv")
add_table(
    ["论文", "研究问题", "优势", "可借鉴点"],
    [[truncate(r["论文"], 18), truncate(r["研究问题"], 38), truncate(r["优势"], 35), truncate(r["可借鉴点"], 36)] for r in key_rows[:7]],
    [1600, 2900, 2400, 2460],
    font_size=7.2,
)

heading("4.7.2 数据集与评价指标", level=3)
paragraph("数据集选择必须服务于病灶分割任务。IDRiD、DDR、FGADR 可作为主要分割数据；APTOS、EyePACS、Messidor 更适合分类预训练、图像质量评估或泛化讨论。评价指标应同时覆盖 mask 重叠、小病灶召回、临床筛查特异性和推理效率。")
add_picture(FIG / "dataset_metric_matrix.png", "图3 数据集与指标选择逻辑")
datasets = read_csv("dataset_comparison.csv")
add_table(
    ["数据集", "任务属性", "适用性", "注意事项"],
    [[truncate(r["数据集"], 18), truncate(r["任务属性"], 25), truncate(r["适用性"], 40), truncate(r["注意事项"], 36)] for r in datasets],
    [1400, 2150, 3300, 2510],
    font_size=7.4,
)
metrics = read_csv("metrics_comparison.csv")
add_table(
    ["指标", "含义", "适用场景", "注意事项"],
    [[truncate(r["指标"], 16), truncate(r["定义/含义"], 38), truncate(r["适用场景"], 28), truncate(r["注意事项"], 34)] for r in metrics[:7]],
    [1300, 3300, 2100, 2660],
    font_size=7.3,
)

heading("4.7.3 存在问题与解决方案", level=3)
add_picture(FIG / "problem_solution_map.png", "图4 研究不足与方案模块映射")
numbered_items([
    "领域适配不足：SAM 来源于自然图像，缺少眼底血管、视盘、病灶形态和 DR 类别先验。解决思路是在 SAM/MedSAM 上加入眼底专用 adapter、LoRA 或 prompt 生成器。",
    "低质量图像影响分割：真实筛查中常见模糊、低对比、过曝和低分辨率。解决思路是加入质量感知增强模块 QEM，并用结构保持损失避免增强过程破坏病灶纹理。",
    "小病灶和类别不平衡：微动脉瘤、软性渗出等目标像素少，整体 Dice 容易掩盖漏检。解决思路是加入小病灶重加权、Boundary/Hausdorff 约束和 lesion-level recall。",
    "推理成本高：SAM/MedSAM 主干较重，不适合筛查部署。解决思路是用 MobileSAM、EdgeSAM 或 KD-SAM 思路进行教师-学生蒸馏。",
])

heading("4.7.4 研究方案与实验规划", level=3)
paragraph("最终方案命名为 LQ-Fundus-SAM，包含 QEM（质量感知增强）、FSA（眼底专用 SAM 适配）、MPP（医学先验 Prompt）、SBC（小病灶/边界约束）和 LSD（轻量化蒸馏）五个模块。总损失函数设计为 L = L_seg + λ1 L_boundary + λ2 L_small + λ3 L_quality + λ4 L_distill。")
add_picture(FIG / "experiment_flow.png", "图5 实验规划流程")
numbered_items([
    "数据集：以 IDRiD 为主数据集，DDR 和 FGADR 用于外部验证或联合训练，APTOS/EyePACS/Messidor 用于预训练或筛查泛化讨论。",
    "Baseline：U-Net、DeepLabv3+、HEDNet+cGAN、RTNet、SAM zero-shot、MedSAM、SAM-Med2D、Medical SAM Adapter，以及 GlanceSeg/TP-DRSeg 思路。",
    "训练设置：冻结 SAM 主干，优先训练 adapter、prompt 生成器、mask head 和增强模块；显存不足时采用 512×512 输入、梯度累积和轻量学生模型。",
    "消融实验：依次去掉 QEM、adapter、文本/类别 prompt、结构先验、小病灶损失和蒸馏模块，验证各模块贡献。",
    "可视化：展示原图、增强图、ground truth、baseline 预测和拟提出模型预测，并单独分析低质量样本、小病灶样本和边界复杂样本。",
])

heading("4.8最终报告与PPT检查", level=2)
numbered_items([
    "Word 报告：最终 DOCX 由 python-docx 构建，并使用 render_docx.py/LibreOffice 渲染页面 PNG 检查版式。",
    "PPT 汇报：最终 PPTX 共 12 页，覆盖背景、论文分析、SOTA、问题、方案、实验规划和心得体会。",
    "合规文件：outputs/compliance_matrix.md 逐条对应实验目的、实验内容、实验步骤和实验要求。",
    "诚信说明：本任务不包含真实训练，因此报告只提出可验证假设，不编造实验结果或性能提升数值。",
])

heading("5.对比")
heading("5.1 开发方式对比", level=2)
paragraph("阶段一的提示词方式适合快速搭建报告框架，但容易把论文摘要堆叠在一起；阶段二的 Agent + 多 Skill 方式把任务拆成可追踪的证据链、表格、图表、方案和交付物，更适合课程实验报告和科研调研。")

heading("5.2 调研效率对比", level=2)
paragraph("提示词方式在早期速度更快，但后续需要大量人工核验。多 Skill 协同模式前期准备较多，却能自动沉淀索引、检索记录、合规矩阵和过程记录，后期修改报告和 PPT 更高效。")

heading("5.3 内容质量对比", level=2)
paragraph("提示词方式容易出现泛泛综述；多 Skill 协同模式可以把“本地 PDF 证据、在线论文、数据集、指标、问题和方案”统一到同一条主线，使报告更像完整实验过程，而不是单纯文献摘抄。")

heading("5.4 功能完整度对比", level=2)
paragraph("阶段二最终产物不仅包括 Word 报告和 PPT，还包括论文索引、检索记录、SOTA/数据集/指标表、图表、合规矩阵、AI 环境与 Skill 调用轨迹、used_skills_md 审批材料，完整度明显高于单一报告文件。")

heading("5.5 核心结论", level=2)
paragraph("对于医学图像技术调研，Agent + 多 Skill 协同模式更适合处理复杂论文库和多阶段交付要求。它能降低信息遗漏和不确定事实被误写的风险，也能把 AI 使用过程本身记录为可审查证据。")

heading("6.心得体会")
paragraph("通过本次实验，我对医学图像分割和 SAM 类基础模型有了更系统的理解。糖尿病视网膜病变筛查不仅需要图像级分类，更需要像素级病灶分割来解释病灶位置、面积和进展。SAM 的 promptable segmentation 范式很有潜力，但不能简单照搬到眼底图像；低质量、小病灶、类别不平衡和临床可解释性都要求模型加入眼底领域先验。")
paragraph("本次实验也让我体会到，AI 辅助调研不能停留在“让模型写一篇综述”。更可靠的做法是先建立论文索引，再逐篇提取证据，随后进行 SOTA、数据集、指标、问题和方案的结构化归纳。Agent 和 Skill 的价值在于把这些环节拆开并留下过程文件，使报告可以复查、可以追溯、可以继续迭代。")
paragraph("最后，我认识到科研调研必须保持诚实边界。本报告提出了 LQ-Fundus-SAM 的研究方案和实验规划，但尚未真实训练模型，因此不能虚构 Dice、IoU、AUC 等数值结果。真正有价值的结论应来自可复现实验，而不是为了让文档看起来完整而编造结果。")

heading("7.参考文献")
refs = [
    "Kirillov A, Mintun E, Ravi N, et al. Segment Anything. arXiv:2304.02643, 2023.",
    "Ma J, He Y, Li F, et al. Segment Anything in Medical Images. arXiv:2304.12306, 2023.",
    "Wu J, Ji W, Liu Y, et al. Medical SAM Adapter: Adapting Segment Anything Model for Medical Image Segmentation. arXiv:2304.12620, 2023.",
    "Cheng J, Ye J, Deng Z, et al. SAM-Med2D. arXiv:2308.16184, 2023.",
    "Ravi N, Gabeur V, Hu Y T, et al. SAM 2: Segment Anything in Images and Videos. arXiv:2408.00714, 2024.",
    "Li W, Xiong X, Xia P, Ju L, Ge Z. TP-DRSeg: Improving Diabetic Retinopathy Lesion Segmentation with Explicit Text-Prompts Assisted SAM. arXiv:2406.15764, 2024.",
    "Huang S, Li J, Xiao Y, Shen N, Xu T. RTNet: Relation Transformer Network for Diabetic Retinopathy Multi-lesion Segmentation. arXiv:2201.11037, 2022.",
    "Jiang H, Gao M, Liu Z, et al. GlanceSeg: Real-time microaneurysm lesion segmentation with gaze-map-guided foundation model for early detection of diabetic retinopathy. arXiv:2311.08075, 2023.",
    "Xiao Q, et al. Improving Lesion Segmentation for Diabetic Retinopathy using Adversarial Learning. arXiv:2007.13854, 2020.",
    "Xia X, Zhan K, Fang Y, Jiang W, Shen F. Lesion-aware network for diabetic retinopathy diagnosis. International Journal of Imaging Systems and Technology, DOI:10.1002/ima.22933.",
    "Wen Y, Luo B, Shi W, et al. SAT-Net: Structure-Aware Transformer-Based Attention Fusion Network for Low-Quality Retinal Fundus Images Enhancement. IEEE Transactions on Multimedia, DOI:10.1109/TMM.2025.3565935.",
    "Patil K D, Palani G, Krishnamurthi G. Efficient Knowledge Distillation of SAM for Medical Image Segmentation. arXiv:2501.16740, 2025.",
    "Zhou C, Li X, Loy C C, Dai B. EdgeSAM: Prompt-In-the-Loop Distillation for SAM. arXiv:2312.06660, 2023.",
    "Moglia A, Leccardi M, Cavicchioli M, et al. Generalist Models in Medical Image Segmentation: A Survey and Performance Comparison with Task-Specific Approaches. arXiv:2506.10825, 2025.",
    "Dai L, Wu L, Li H, et al. A deep learning system for detecting diabetic retinopathy across the disease spectrum. Nature Communications, DOI:10.1038/s41467-021-23458-5.",
    "Dai L, Sheng B, Chen T, et al. A deep learning system for predicting time to progression of diabetic retinopathy. Nature Medicine, DOI:10.1038/s41591-023-02702-z.",
    "Li J, Guan Z, Wang J, et al. Integrated image-based deep learning and language models for primary diabetes care. Nature Medicine, DOI:10.1038/s41591-024-03139-8.",
]
for i, ref in enumerate(refs, 1):
    paragraph(f"{i}. {ref}", size=9.5, first_line=False, line=1.15, after=2)

doc.core_properties.title = "AI开发技术-实验4.1-医学图像技术调研报告"
doc.core_properties.author = "胡文韬"
doc.save(FINAL_DOCX)
print(FINAL_DOCX)
